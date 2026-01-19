"""
Bitbucket Release Documentation Generator
Generates release documentation DOCX files from Bitbucket repository commits and tags.
"""

from atlassian import Bitbucket
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import List, Dict, Tuple
import os


class BitbucketReleaseDocumentGenerator:
    """Generate release documents from Bitbucket repository data."""
    
    def __init__(self, server_url: str, username: str, password: str, 
                 project_key: str, repo_slug: str):
        """
        Initialize Bitbucket connection and document generator.
        
        Args:
            server_url: Bitbucket server URL (e.g., 'https://bitbucket.company.com')
            username: Bitbucket username
            password: Bitbucket password/token
            project_key: Project key in Bitbucket (e.g., 'PROJ')
            repo_slug: Repository slug name
        """
        self.bitbucket = Bitbucket(url=server_url, username=username, password=password)
        self.project_key = project_key
        self.repo_slug = repo_slug
        self.server_url = server_url
        
    def get_prod_tag_commit(self) -> Dict:
        """
        Find the commit tagged with 'prod-server'.
        
        Returns:
            Dict containing commit information
        """
        print("Fetching prod-server tag...")
        tags = self.bitbucket.get_repo_tags(self.project_key, self.repo_slug)
        
        for tag in tags:
            if 'prod-server' in tag.get('displayId', ''):
                commit_id = tag.get('latestCommit')
                return self.bitbucket.get_commit(
                    self.project_key, 
                    self.repo_slug, 
                    commit_id
                )
        
        raise ValueError("prod-server tag not found in repository")
    
    def get_commits_after_tag(self, since_commit_id: str) -> List[Dict]:
        """
        Get all commits in master branch after the prod-server tag commit.
        
        Args:
            since_commit_id: Commit ID to start from
            
        Returns:
            List of commit dictionaries
        """
        print("Fetching commits after prod-server tag...")
        commits = []
        
        # Get commits from master branch after the specified commit
        response = self.bitbucket.get_commits(
            self.project_key,
            self.repo_slug,
            all=True
        )
        
        # Filter commits that are newer than the prod-server tag commit
        found_marker = False
        for commit in response:
            if commit.get('id') == since_commit_id:
                found_marker = True
                break
            commits.append(commit)
        
        print(f"Found {len(commits)} commits after prod-server tag")
        return commits
    
    def get_tags_for_commit(self, commit_id: str) -> List[str]:
        """
        Get all tags for a specific commit that contain 'prod-server'.
        
        Args:
            commit_id: Commit ID
            
        Returns:
            List of tag names
        """
        print("Fetching tags for commit...")
        try:
            tags = self.bitbucket.get_commit_tags(
                self.project_key,
                self.repo_slug,
                commit_id
            )
            # Filter tags containing 'prod-server'
            prod_tags = [tag.get('displayId') for tag in tags 
                        if 'prod-server' in tag.get('displayId', '')]
            return prod_tags
        except Exception as e:
            print(f"Note: Could not fetch tags: {e}")
            return []
    
    def get_changed_files_in_deployment(self, commits: List[Dict]) -> List[Dict]:
        """
        Get all files changed in the deployment folder across commits.
        
        Args:
            commits: List of commit dictionaries
            
        Returns:
            List of changed files in deployment folder
        """
        print("Fetching changed files in deployment folder...")
        deployment_files = {}
        
        for commit in commits:
            commit_id = commit.get('id')
            try:
                changes = self.bitbucket.get_commit_changes(
                    self.project_key,
                    self.repo_slug,
                    commit_id
                )
                
                for change in changes:
                    file_path = change.get('path', {}).get('toString', '')
                    if 'deployment' in file_path.lower():
                        if file_path not in deployment_files:
                            deployment_files[file_path] = {
                                'path': file_path,
                                'type': change.get('type', 'UNKNOWN'),
                                'commits': []
                            }
                        deployment_files[file_path]['commits'].append(commit_id[:7])
            except Exception as e:
                print(f"Warning: Could not fetch changes for commit {commit_id}: {e}")
        
        print(f"Found {len(deployment_files)} deployment files changed")
        return list(deployment_files.values())
    
    def get_commit_url(self, commit_id: str) -> str:
        """Generate commit URL."""
        return f"{self.server_url}/projects/{self.project_key}/repos/{self.repo_slug}/commits/{commit_id}"
    
    # ===== Document Creation Methods =====
    
    def create_implementation_plan(self, deployment_files: List[Dict], 
                                   input_text: str = "") -> Document:
        """
        Create Implementation Plan document.
        
        Args:
            deployment_files: List of changed files
            input_text: User input text for implementation details
            
        Returns:
            Document object
        """
        doc = Document()
        doc.add_heading('Implementation Plan - Change Request', 0)
        
        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Project: {self.project_key} | Repository: {self.repo_slug}")
        doc.add_paragraph()
        
        # Implementation Overview
        doc.add_heading('Implementation Overview', level=1)
        doc.add_paragraph(input_text if input_text else "Please provide implementation details")
        doc.add_paragraph()
        
        # Files Changed in Deployment
        doc.add_heading('Files Changed in Deployment Folder', level=1)
        
        if deployment_files:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'File Path'
            header_cells[1].text = 'Type'
            header_cells[2].text = 'Commits'
            
            # Data rows
            for file_info in deployment_files:
                row_cells = table.add_row().cells
                row_cells[0].text = file_info['path']
                row_cells[1].text = file_info['type']
                row_cells[2].text = ', '.join(file_info['commits'])
        else:
            doc.add_paragraph("No deployment files changed in this release")
        
        doc.add_paragraph()
        doc.add_paragraph("---")
        doc.add_paragraph("Change Type: Release Deployment")
        doc.add_paragraph("Status: Pending Review")
        
        return doc
    
    def create_pre_test_plan(self) -> Document:
        """Create PRE Test Plan document."""
        doc = Document()
        doc.add_heading('PRE Test Plan - Change Request', 0)
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Pre-Deployment Testing', level=1)
        doc.add_paragraph("Test Scope:")
        doc.add_paragraph("• Code quality checks", style='List Bullet')
        doc.add_paragraph("• Unit tests verification", style='List Bullet')
        doc.add_paragraph("• Integration tests", style='List Bullet')
        doc.add_paragraph("• Security scanning", style='List Bullet')
        
        doc.add_paragraph()
        doc.add_heading('Test Results', level=1)
        doc.add_paragraph("All tests must pass before proceeding to deployment.")
        
        # Add checkbox table for test results
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Test Type'
        header_cells[1].text = 'Status'
        header_cells[2].text = 'Notes'
        
        test_types = ['Code Quality', 'Unit Tests', 'Integration Tests', 'Security Scan']
        for idx, test_type in enumerate(test_types, 1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = test_type
            row_cells[1].text = '[ ] Pass  [ ] Fail'
            row_cells[2].text = ''
        
        return doc
    
    def create_post_test_plan(self) -> Document:
        """Create POST Test Plan document."""
        doc = Document()
        doc.add_heading('POST Test Plan - Change Request', 0)
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Post-Deployment Validation', level=1)
        doc.add_paragraph("Validation Steps:")
        doc.add_paragraph("• System health check", style='List Bullet')
        doc.add_paragraph("• Service availability verification", style='List Bullet')
        doc.add_paragraph("• Database consistency check", style='List Bullet')
        doc.add_paragraph("• Application logs review", style='List Bullet')
        doc.add_paragraph("• User acceptance testing", style='List Bullet')
        
        doc.add_paragraph()
        doc.add_heading('Success Criteria', level=1)
        doc.add_paragraph("All deployment targets must be operational with no critical errors in logs.")
        
        # Validation checklist
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Validation Item'
        header_cells[1].text = 'Status'
        
        items = ['Health Check', 'Services Running', 'Database Sync', 'Error Log Check', 'UAT Approval']
        for idx, item in enumerate(items, 1):
            row_cells = table.rows[idx].cells
            row_cells[0].text = item
            row_cells[1].text = '[ ] OK  [ ] Failed'
        
        return doc
    
    def create_rollback_plan(self, tags: List[str]) -> Document:
        """
        Create Rollback Plan document.
        
        Args:
            tags: List of tags associated with the release
            
        Returns:
            Document object
        """
        doc = Document()
        doc.add_heading('Rollback Plan - Change Request', 0)
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Release Information', level=1)
        if tags:
            doc.add_paragraph(f"Release Tags: {', '.join(tags)}")
        else:
            doc.add_paragraph("Release Tags: Not available")
        doc.add_paragraph()
        
        doc.add_heading('Rollback Strategy', level=1)
        doc.add_paragraph("Variable text placeholder for rollback details:")
        doc.add_paragraph("[Insert rollback strategy details here]")
        doc.add_paragraph()
        
        doc.add_heading('Rollback Procedures', level=1)
        doc.add_paragraph("Step 1: Notify stakeholders of rollback decision", style='List Number')
        doc.add_paragraph("Step 2: Take application offline for maintenance", style='List Number')
        doc.add_paragraph("Step 3: Restore database from backup", style='List Number')
        doc.add_paragraph("Step 4: Revert application code to previous release", style='List Number')
        doc.add_paragraph("Step 5: Verify system functionality", style='List Number')
        doc.add_paragraph("Step 6: Bring application back online", style='List Number')
        
        doc.add_paragraph()
        doc.add_heading('Rollback Triggers', level=1)
        doc.add_paragraph("Rollback should be initiated if:")
        doc.add_paragraph("• Critical application error occurs", style='List Bullet')
        doc.add_paragraph("• Database corruption is detected", style='List Bullet')
        doc.add_paragraph("• Service availability drops below SLA", style='List Bullet')
        doc.add_paragraph("• Data integrity issues are found", style='List Bullet')
        
        return doc
    
    def create_code_change_review(self, commits: List[Dict]) -> Document:
        """
        Create Code Change Review document with commit links.
        
        Args:
            commits: List of commit dictionaries
            
        Returns:
            Document object
        """
        doc = Document()
        doc.add_heading('Code Change Review - Change Request', 0)
        
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Repository: {self.project_key}/{self.repo_slug}")
        doc.add_paragraph()
        
        doc.add_heading('Commits Included in Release', level=1)
        
        if commits:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Commit ID'
            header_cells[1].text = 'Author'
            header_cells[2].text = 'Date'
            header_cells[3].text = 'Message'
            
            # Add commits to table
            for commit in commits:
                row_cells = table.add_row().cells
                
                commit_id = commit.get('id', '')
                commit_short = commit_id[:7] if len(commit_id) > 7 else commit_id
                
                row_cells[0].text = commit_short
                row_cells[1].text = commit.get('author', {}).get('name', 'Unknown')
                row_cells[2].text = commit.get('authorTimestamp', '')
                row_cells[3].text = commit.get('message', '')[:100]
        else:
            doc.add_paragraph("No commits found in this release")
        
        doc.add_paragraph()
        doc.add_heading('Commit Details with Links', level=1)
        
        if commits:
            for commit in commits:
                commit_id = commit.get('id', '')
                commit_short = commit_id[:7] if len(commit_id) > 7 else commit_id
                
                # Add hyperlink using run
                paragraph = doc.add_paragraph(style='List Bullet')
                run = paragraph.add_run(f"{commit_short}: ")
                
                commit_url = self.get_commit_url(commit_id)
                run.hyperlink.href = commit_url
                
                message = commit.get('message', '').split('\n')[0]
                paragraph.add_run(message)
        
        doc.add_paragraph()
        doc.add_heading('Code Review Checklist', level=1)
        
        checklist_items = [
            'All changes have been reviewed',
            'Code follows standards and best practices',
            'No security vulnerabilities identified',
            'Tests are included and passing',
            'Documentation is updated',
            'Performance impact is acceptable'
        ]
        
        for item in checklist_items:
            doc.add_paragraph(f"[ ] {item}", style='List Bullet')
        
        return doc
    
    def save_documents(self, output_dir: str = ".", implementation_text: str = "", 
                      rollback_text: str = "") -> Dict[str, str]:
        """
        Generate and save all required DOCX files.
        
        Args:
            output_dir: Directory to save documents
            implementation_text: Text for implementation plan
            rollback_text: Text for rollback plan
            
        Returns:
            Dictionary with file names and their paths
        """
        print("\n=== Starting Document Generation ===\n")
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        # Get data from Bitbucket
        prod_tag_commit = self.get_prod_tag_commit()
        prod_commit_id = prod_tag_commit.get('id')
        
        commits = self.get_commits_after_tag(prod_commit_id)
        tags = self.get_tags_for_commit(prod_commit_id)
        deployment_files = self.get_changed_files_in_deployment(commits)
        
        # Generate documents
        docs = {
            'Implementation_plan_CHG.docx': self.create_implementation_plan(
                deployment_files, 
                implementation_text
            ),
            'PRE_test_plan_CHG.docx': self.create_pre_test_plan(),
            'POST_test_plan_CHG.docx': self.create_post_test_plan(),
            'Rollback_plan_CHG.docx': self.create_rollback_plan(tags),
            'Code_change_Review_CHG.docx': self.create_code_change_review(commits)
        }
        
        # Save documents
        saved_files = {}
        for filename, doc in docs.items():
            filepath = os.path.join(output_dir, filename)
            doc.save(filepath)
            saved_files[filename] = filepath
            print(f"✓ Created: {filepath}")
        
        print(f"\n=== Document Generation Complete ===")
        print(f"Documents saved to: {os.path.abspath(output_dir)}")
        
        return saved_files


def main():
    """Main function to run the document generator."""
    
    # Configuration - Update these with your Bitbucket details
    BITBUCKET_URL = "https://bitbucket.company.com"  # Change to your Bitbucket server URL
    USERNAME = "your_username"  # Change to your username
    PASSWORD = "your_password_or_token"  # Change to your password or API token
    PROJECT_KEY = "PROJ"  # Change to your project key
    REPO_SLUG = "repository-name"  # Change to your repository slug
    
    # Input texts for documents
    IMPLEMENTATION_TEXT = """
    This release includes critical bug fixes and performance improvements.
    
    Key Changes:
    - Fixed authentication issues
    - Improved database query performance
    - Enhanced error handling
    """
    
    ROLLBACK_TEXT = """
    In case of deployment issues, execute the following rollback plan:
    1. Revert to the previous stable release tag
    2. Run database migration scripts in reverse
    3. Clear application cache
    4. Verify all services are operational
    """
    
    # Create generator
    generator = BitbucketReleaseDocumentGenerator(
        server_url=BITBUCKET_URL,
        username=USERNAME,
        password=PASSWORD,
        project_key=PROJECT_KEY,
        repo_slug=REPO_SLUG
    )
    
    # Generate and save documents
    saved_files = generator.save_documents(
        output_dir="./release_documents",
        implementation_text=IMPLEMENTATION_TEXT,
        rollback_text=ROLLBACK_TEXT
    )
    
    print("\nGenerated Files:")
    for filename, filepath in saved_files.items():
        print(f"  - {filename}")


if __name__ == "__main__":
    main()
